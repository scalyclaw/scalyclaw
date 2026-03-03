import sys
import json
import os
import csv
import io
import re


VALID_FROM = ("docx", "xlsx", "html", "json", "yaml", "markdown", "csv")
VALID_TO = ("markdown", "csv", "json", "yaml", "html", "text", "docx")

FORMAT_EXTENSIONS = {
    "markdown": "md",
    "csv": "csv",
    "json": "json",
    "yaml": "yaml",
    "html": "html",
    "text": "txt",
    "docx": "docx",
}


def get_workspace():
    return os.environ.get("WORKSPACE_DIR", ".")


def read_input(data):
    """Read content from file_path or inline content."""
    file_path = data.get("file_path")
    content = data.get("content")

    if file_path:
        if not os.path.isfile(file_path):
            return None, f"File not found: {file_path}"
        return file_path, None
    elif content is not None:
        return content, None
    else:
        return None, "Either file_path or content must be provided"


def write_output(content_str, output_filename, to_format):
    """Write content to a file in WORKSPACE_DIR."""
    workspace = get_workspace()
    if not output_filename:
        ext = FORMAT_EXTENSIONS.get(to_format, to_format)
        output_filename = f"converted.{ext}"
    output_path = os.path.join(workspace, output_filename)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content_str)
    return output_path


def write_output_binary(content_bytes, output_filename, to_format):
    """Write binary content to a file in WORKSPACE_DIR."""
    workspace = get_workspace()
    if not output_filename:
        ext = FORMAT_EXTENSIONS.get(to_format, to_format)
        output_filename = f"converted.{ext}"
    output_path = os.path.join(workspace, output_filename)
    with open(output_path, "wb") as f:
        f.write(content_bytes)
    return output_path


# --- Conversion functions ---


def docx_to_markdown(file_path):
    """Convert DOCX to Markdown."""
    from docx import Document

    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text

        # Handle headings
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            level = min(level, 6)
            lines.append(f"{'#' * level} {text}")
        elif style_name.startswith("List"):
            lines.append(f"- {text}")
        else:
            # Handle inline formatting from runs
            formatted_parts = []
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                formatted_parts.append(run_text)

            if formatted_parts:
                lines.append("".join(formatted_parts))
            elif text:
                lines.append(text)
            else:
                lines.append("")

    return "\n\n".join(lines)


def markdown_to_docx(content_str):
    """Convert Markdown to DOCX. Returns bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    lines = content_str.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Unordered list items
        list_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if list_match:
            text = list_match.group(1)
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(para, text)
            i += 1
            continue

        # Ordered list items
        ordered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered_match:
            text = ordered_match.group(1)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_runs(para, text)
            i += 1
            continue

        # Code blocks
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code_text = "\n".join(code_lines)
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        _add_formatted_runs(para, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_runs(paragraph, text):
    """Parse inline markdown formatting and add runs to paragraph."""
    from docx.shared import Pt

    # Pattern to match bold+italic, bold, italic, inline code
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"   # bold+italic
        r"|(\*\*(.+?)\*\*)"      # bold
        r"|(\*(.+?)\*)"          # italic
        r"|(`(.+?)`)"            # inline code
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add preceding plain text
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # bold+italic
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):  # bold
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6):  # italic
            run = paragraph.add_run(match.group(6))
            run.italic = True
        elif match.group(8):  # inline code
            run = paragraph.add_run(match.group(8))
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        last_end = match.end()

    # Add remaining plain text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def xlsx_to_csv(file_path):
    """Convert XLSX to CSV string."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    ws = wb.active

    output = io.StringIO()
    writer = csv.writer(output)

    for row in ws.iter_rows(values_only=True):
        writer.writerow(
            [str(cell) if cell is not None else "" for cell in row]
        )

    wb.close()
    return output.getvalue()


def html_to_markdown_convert(content_str):
    """Convert HTML to Markdown."""
    from markdownify import markdownify

    return markdownify(content_str, heading_style="ATX").strip()


def html_to_text_convert(content_str):
    """Convert HTML to plain text."""
    from markdownify import markdownify

    md = markdownify(content_str, heading_style="ATX")
    # Strip markdown formatting for plain text
    text = re.sub(r"[#*_`\[\]]", "", md)
    text = re.sub(r"\(https?://[^\)]+\)", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def json_to_yaml_convert(content_str):
    """Convert JSON to YAML."""
    import yaml

    data = json.loads(content_str)
    return yaml.dump(data, default_flow_style=False, allow_unicode=True)


def yaml_to_json_convert(content_str):
    """Convert YAML to JSON."""
    import yaml

    data = yaml.safe_load(content_str)
    return json.dumps(data, indent=2, ensure_ascii=False)


def csv_to_json_convert(content_str):
    """Convert CSV to JSON array using DictReader."""
    reader = csv.DictReader(io.StringIO(content_str))
    rows = list(reader)
    return json.dumps(rows, indent=2, ensure_ascii=False)


def json_to_csv_convert(content_str):
    """Convert JSON array to CSV."""
    data = json.loads(content_str)
    if not isinstance(data, list) or len(data) == 0:
        raise ValueError("JSON must be an array of objects for CSV conversion")

    output = io.StringIO()
    fieldnames = list(data[0].keys())
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    for row in data:
        writer.writerow(row)
    return output.getvalue()


def markdown_to_html_convert(content_str):
    """Convert Markdown to HTML."""
    import markdown

    return markdown.markdown(content_str, extensions=["tables", "fenced_code"])


def main():
    try:
        data = json.loads(sys.stdin.read())
        from_format = data.get("from_format", "").lower()
        to_format = data.get("to_format", "").lower()
        output_filename = data.get("output_filename")

        if not from_format:
            print(json.dumps({"error": "Missing required field: from_format"}))
            return
        if not to_format:
            print(json.dumps({"error": "Missing required field: to_format"}))
            return
        if from_format not in VALID_FROM:
            print(
                json.dumps(
                    {
                        "error": f"Invalid from_format: {from_format}. "
                        f"Must be one of: {', '.join(VALID_FROM)}"
                    }
                )
            )
            return
        if to_format not in VALID_TO:
            print(
                json.dumps(
                    {
                        "error": f"Invalid to_format: {to_format}. "
                        f"Must be one of: {', '.join(VALID_TO)}"
                    }
                )
            )
            return

        input_data, err = read_input(data)
        if err:
            print(json.dumps({"error": err}))
            return

        # Determine if input_data is a file path or inline content
        is_file = isinstance(input_data, str) and os.path.isfile(input_data)

        sys.stderr.write(f"Converting {from_format} -> {to_format}\n")

        converted = None
        is_binary = False

        # --- DOCX conversions ---
        if from_format == "docx" and to_format == "markdown":
            if not is_file:
                print(json.dumps({"error": "DOCX conversion requires a file_path"}))
                return
            converted = docx_to_markdown(input_data)

        elif from_format == "docx" and to_format == "text":
            if not is_file:
                print(json.dumps({"error": "DOCX conversion requires a file_path"}))
                return
            from docx import Document

            doc = Document(input_data)
            converted = "\n".join(para.text for para in doc.paragraphs)

        elif from_format == "markdown" and to_format == "docx":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = markdown_to_docx(content_str)
            is_binary = True

        # --- XLSX conversions ---
        elif from_format == "xlsx" and to_format == "csv":
            if not is_file:
                print(json.dumps({"error": "XLSX conversion requires a file_path"}))
                return
            converted = xlsx_to_csv(input_data)

        elif from_format == "xlsx" and to_format == "json":
            if not is_file:
                print(json.dumps({"error": "XLSX conversion requires a file_path"}))
                return
            csv_str = xlsx_to_csv(input_data)
            converted = csv_to_json_convert(csv_str)

        # --- HTML conversions ---
        elif from_format == "html" and to_format == "markdown":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = html_to_markdown_convert(content_str)

        elif from_format == "html" and to_format == "text":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = html_to_text_convert(content_str)

        # --- Markdown conversions ---
        elif from_format == "markdown" and to_format == "html":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = markdown_to_html_convert(content_str)

        # --- JSON/YAML conversions ---
        elif from_format == "json" and to_format == "yaml":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = json_to_yaml_convert(content_str)

        elif from_format == "yaml" and to_format == "json":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = yaml_to_json_convert(content_str)

        # --- CSV conversions ---
        elif from_format == "csv" and to_format == "json":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = csv_to_json_convert(content_str)

        elif from_format == "json" and to_format == "csv":
            content_str = (
                open(input_data, "r", encoding="utf-8").read()
                if is_file
                else input_data
            )
            converted = json_to_csv_convert(content_str)

        else:
            print(
                json.dumps(
                    {
                        "error": f"Unsupported conversion: {from_format} -> {to_format}"
                    }
                )
            )
            return

        # Build output
        result = {
            "from_format": from_format,
            "to_format": to_format,
        }

        # Binary outputs (DOCX) always go to file
        if is_binary:
            out_path = write_output_binary(converted, output_filename, to_format)
            result["file_path"] = out_path
            sys.stderr.write(f"Output written to: {out_path}\n")
        else:
            # For text-based outputs: write to file and also include content
            out_path = write_output(converted, output_filename, to_format)
            result["file_path"] = out_path
            result["content"] = converted
            sys.stderr.write(f"Output written to: {out_path}\n")

        print(json.dumps(result))
    except Exception as e:
        sys.stderr.write(f"Error: {e}\n")
        print(json.dumps({"error": str(e)}))


if __name__ == "__main__":
    main()
